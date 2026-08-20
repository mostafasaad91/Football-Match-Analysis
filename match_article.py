"""The match article: one publishable read, built from the fixture's own numbers.

The report is a reference — every visual, every metric, a paragraph under each.
An article is not that. It has to pick an argument, order the evidence behind
it, and stop. So this does not walk the visuals and describe them; it derives a
small set of *findings* from the frames, ranks them by how far apart the two
sides actually were, and gives the strongest five or six a section each with
the visuals that evidence them.

Two rules hold everywhere in here, and both come from defects the report
shipped:

- a sentence names whichever side the numbers name, never a fixed one. Four
  readings in the PDF said "the away side" and meant "the leader", so they
  contradicted the figures printed beside them whenever the home team led.
- a number is followed by what it cost or bought, not by a restatement. "70.8%
  field tilt" is a measurement; "seventy per cent of the match in the other
  half, for 1.08 xG" is a finding.

``build_article`` returns structured data. ``render_docx`` writes it. Keeping
those apart is what lets the same argument reach a Word file for publishing and
the report's own pages without the two drifting.
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd

from frame_values import (number as _number, ratio as _ratio, text as _text,
                          whole as _whole)

# The argued sections have no ceiling: every finding the match supports is
# worth printing, and the appendix behind them carries a reading under every
# remaining board. The floor stays, because a piece shorter than this is a
# summary rather than a read.
TARGET_WORDS = (1200, None)

# Player radars carried per side, matching the report's own appendix.
RADARS_PER_TEAM = 5


# --------------------------------------------------------------------------
# reading the frames
# --------------------------------------------------------------------------

def _num(row, key, default=0.0) -> float:
    try:
        value = float(row.get(key, default))
    except (TypeError, ValueError):
        return float(default)
    return float(default) if pd.isna(value) else value


def _side(team_metrics: pd.DataFrame, side: str) -> pd.Series:
    rows = team_metrics[team_metrics["side"].astype(str).eq(side)]
    return rows.iloc[0] if not rows.empty else pd.Series(dtype=float)


def _xg_row(xg: pd.DataFrame, team: str) -> pd.Series:
    rows = xg[xg["team"].astype(str).str.lower().eq(str(team).lower())]
    return rows.iloc[0] if not rows.empty else pd.Series(dtype=float)


def _slug(name: str) -> str:
    import re

    return re.sub(r"[^a-z0-9]+", "_", str(name).lower()).strip("_") or "team"


def _plural(count: float, one: str, many: str | None = None) -> str:
    return one if int(round(count)) == 1 else (many or one + "s")


def _spell(value: float) -> str:
    """Small integers read better as words in running prose."""
    words = {0: "no", 1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
             6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten",
             11: "eleven", 12: "twelve"}
    number = int(round(value))
    return words.get(number, str(number))


# --------------------------------------------------------------------------
# structure
# --------------------------------------------------------------------------

@dataclass
class Section:
    heading: str
    paragraphs: list[str]
    visuals: list[Path] = field(default_factory=list)
    pull_quote: str | None = None
    # True for the sections whose job is to show rather than to argue: the
    # player profiles and the closing gallery. The argument's own sections stay
    # tight, and identifying them by position broke the moment the profiles
    # stopped being last.
    gallery: bool = False

    def words(self) -> int:
        return sum(len(p.split()) for p in self.paragraphs)


@dataclass
class Article:
    title: str
    standfirst: str
    strap: str
    sections: list[Section]
    cover: Path | None = None
    home: str = ""
    away: str = ""
    # The report's own context, so each visual in the appendix can carry the
    # analytical note the PDF writes under it rather than a bare caption.
    context: dict | None = None

    def words(self) -> int:
        return sum(section.words() for section in self.sections)

    def narrative_words(self) -> int:
        """The argued part, without the appendix that shows the rest."""
        return sum(s.words() for s in self.sections if not s.gallery)


@dataclass
class Finding:
    """One thing the match did, with the size of it and the evidence for it."""

    key: str
    weight: float          # relative gap, for ranking
    section: Section


# --------------------------------------------------------------------------
# the findings
# --------------------------------------------------------------------------

class _Match:
    """Everything the prose needs, read once from the frames."""

    def __init__(self, events, xg, team_metrics, player_metrics, info, out_dir):
        self.out = Path(out_dir)
        self.home = str(info["home_name"])
        self.away = str(info["away_name"])
        self.home_slug, self.away_slug = _slug(self.home), _slug(self.away)
        self.hm, self.am = _side(team_metrics, "home"), _side(team_metrics, "away")
        self.hx, self.ax = _xg_row(xg, self.home), _xg_row(xg, self.away)
        self.players = player_metrics
        self.competition = _text(info.get("competition"))

        self.home_goals = int(_num(self.hx, "goals"))
        self.away_goals = int(_num(self.ax, "goals"))
        self.home_xg, self.away_xg = _num(self.hx, "xG"), _num(self.ax, "xG")
        self.score = f"{self.home_goals}–{self.away_goals}"

        self.first_goal = self._first_goal(events)

        if self.home_goals > self.away_goals:
            self.winner, self.loser = self.home, self.away
        elif self.away_goals > self.home_goals:
            self.winner, self.loser = self.away, self.home
        else:
            self.winner = self.loser = None

    @staticmethod
    def _first_goal(events):
        """(when, scorer) for the opening goal, or None.

        ``when`` is a phrase rather than a number because "minute 0" is not how
        anyone describes a goal scored inside the first minute — and a goal that
        early is exactly the one worth describing precisely, since it is the
        whole reason the two sides never played a level match.
        """
        try:
            goals = events[events["is_goal"].astype(str).str.lower().eq("true")]
            goals = goals.sort_values(["minute", "second"], kind="stable")
            if goals.empty:
                return None
            row = goals.iloc[0]
            minute = int(float(row["minute"]))
            try:
                second = int(float(row.get("second")))
            except (TypeError, ValueError):
                second = 0
            if minute == 0:
                when = f"after {second} seconds" if second else "from the kick-off"
            else:
                when = f"in minute {minute + 1}"
            # `nan or ""` returns nan — a float NaN is truthy — so an
            # unattributed goal produced "through nan".
            return when, _text(row.get("player")), minute
        except Exception:
            return None

    # -- helpers ---------------------------------------------------------
    def lead(self, home_value, away_value, tolerance: float = 0.0):
        """(leader, trailer, level) for a pair, named by the numbers."""
        if abs(home_value - away_value) <= tolerance:
            return self.home, self.away, True
        if home_value >= away_value:
            return self.home, self.away, False
        return self.away, self.home, False

    def gap(self, home_value, away_value) -> float:
        peak = max(abs(home_value), abs(away_value))
        return abs(home_value - away_value) / peak if peak else 0.0

    def of(self, team: str, home_value, away_value):
        return home_value if team == self.home else away_value

    def visual(self, *names: str) -> list[Path]:
        found = []
        for name in names:
            path = self.out / name
            if path.exists():
                found.append(path)
        return found

    def team_visual(self, pattern: str, team: str) -> list[Path]:
        slug = self.home_slug if team == self.home else self.away_slug
        return self.visual(pattern.format(slug=slug))

    def radar(self, team: str, player: str) -> Path | None:
        """The exported radar for one player, if it was written."""
        folder = self.out / "player_radars" / str(team).replace(" ", "_")
        candidate = folder / (str(player).replace(" ", "_") + ".png")
        if candidate.exists():
            return candidate
        # Fall back to a case-insensitive match: the exporter and the ranking
        # agree on the name, but not always on its punctuation.
        if folder.exists():
            wanted = str(player).replace(" ", "_").lower()
            for found in folder.glob("*.png"):
                if found.stem.lower() == wanted:
                    return found
        return None

    def top_player(self, team: str | None = None):
        if self.players is None or self.players.empty:
            return None
        frame = self.players.copy()
        frame["sequence_xT"] = pd.to_numeric(
            frame["sequence_xT"], errors="coerce").fillna(0)
        if team is not None and "team" in frame.columns:
            frame = frame[frame["team"].astype(str).eq(team)]
        if frame.empty:
            return None
        return frame.sort_values("sequence_xT", ascending=False).iloc[0]


def _finding_result(m: _Match) -> Finding | None:
    """The result against the process behind it. Always the opening argument."""
    leader, trailer, level = m.lead(m.home_xg, m.away_xg, tolerance=0.15)
    combined = m.home_xg + m.away_xg
    goals = m.home_goals + m.away_goals

    if m.winner and not level and leader != m.winner:
        heading = f"{m.winner} won the match {m.loser} played better in"
        first = (
            f"{m.winner} took the points and {m.loser} took the chances. "
            f"{m.loser} finished on {m.of(m.loser, m.home_xg, m.away_xg):.2f} expected goals "
            f"against {m.of(m.winner, m.home_xg, m.away_xg):.2f}, and lost anyway. That is not a "
            f"complaint about luck; it is the starting point of the analysis, "
            f"because everything either side did before the final touch is "
            f"described by the first number and everything that happened to the "
            f"scoreboard by the second."
        )
        weight = m.gap(m.home_xg, m.away_xg) + 0.5
    elif m.winner:
        heading = f"{m.winner} deserved it, and the margin says how"
        first = (
            f"{m.winner} won on the scoreboard and won the underlying match with it: "
            f"{m.of(m.winner, m.home_xg, m.away_xg):.2f} expected goals against "
            f"{m.of(m.loser, m.home_xg, m.away_xg):.2f}. When the result and the process "
            f"agree, the interesting question moves from whether the scoreline was "
            f"fair to how the advantage was built, and that is what the rest of this "
            f"is about."
        )
        weight = m.gap(m.home_xg, m.away_xg) + 0.4
    else:
        heading = "A draw that was not an even match"
        first = (
            f"The scoreline finished level at {m.score}. The expected goals did not: "
            f"{m.home} {m.home_xg:.2f}, {m.away} {m.away_xg:.2f}. A draw is the least "
            f"informative result football produces, and the only way to read one is to "
            f"go past it."
        )
        weight = m.gap(m.home_xg, m.away_xg) + 0.3

    second = (
        f"{_spell(goals).capitalize()} {_plural(goals, 'goal')} came from {combined:.2f} "
        f"combined expected goals, so finishing ran "
        f"{'ahead of' if goals > combined else 'behind'} the chances on the night. "
        f"Read that as a warning about the sample, not a verdict on the players: "
        f"conversion is the noisiest thing in the match and the least likely part of "
        f"it to repeat."
    )
    # Which caveat is worth making depends on how far the finishing ran from the
    # chances, so the paragraph is chosen by that distance rather than fixed.
    overperformance = goals - combined
    if overperformance > 0.8:
        third = (
            f"A gap that size between {goals} scored and {combined:.2f} created is not a "
            f"skill the players demonstrated; it is the shortest sample in football "
            f"behaving like one. Over a season the two numbers converge, which is why "
            f"the chances are the part worth arguing from and the conversion is the part "
            f"worth noting and setting aside."
        )
    elif overperformance < -0.8:
        third = (
            f"{combined:.2f} expected goals produced {_spell(goals)} "
            f"{_plural(goals, 'goal')}, so both sides finished below what the chances "
            f"were worth. That is worth naming and not worth over-reading: a match this "
            f"short cannot separate a genuine finishing problem from an ordinary "
            f"afternoon in front of goal, and the shape of the chances is the more "
            f"durable evidence either way."
        )
    else:
        third = (
            f"Conversion tracked the chances closely here — {goals} from "
            f"{combined:.2f} — which is the least common thing a match does and the "
            f"most convenient for reading one. When finishing neither flatters nor "
            f"hides a performance, the underlying numbers can be taken more or less at "
            f"face value, and the argument moves straight to how the chances were built."
        )
    return Finding("result", weight, Section(
        heading, [first, second, third],
        m.visual("01_xg_flow.png", "46_goal_origins.png"),
        pull_quote=f"{m.home} {m.home_xg:.2f} xG  ·  {m.away} {m.away_xg:.2f} xG",
    ))


def _finding_territory(m: _Match) -> Finding | None:
    """Who held the ball and the ground, and what it bought."""
    home_tilt, away_tilt = _num(m.hm, "field_tilt"), _num(m.am, "field_tilt")
    if not (home_tilt or away_tilt):
        return None
    tilt_leader, tilt_trailer, level = m.lead(home_tilt, away_tilt, tolerance=3.0)
    if level:
        return None
    tilt_value = m.of(tilt_leader, home_tilt, away_tilt)
    leader_xg = m.of(tilt_leader, m.home_xg, m.away_xg)
    trailer_xg = m.of(tilt_trailer, m.home_xg, m.away_xg)
    thirds = m.of(tilt_leader, _num(m.hm, "final_third_entries"),
                  _num(m.am, "final_third_entries"))
    boxes = m.of(tilt_leader, _num(m.hm, "box_entries"), _num(m.am, "box_entries"))
    other_boxes = m.of(tilt_trailer, _num(m.hm, "box_entries"),
                       _num(m.am, "box_entries"))

    if leader_xg < trailer_xg:
        heading = "Territory is not the same thing as threat"
        opening = (
            f"{tilt_leader} spent {tilt_value:.1f}% of the completed passing in the "
            f"final third and got {leader_xg:.2f} expected goals for it. "
            f"{tilt_trailer}, with the smaller share, got {trailer_xg:.2f}."
        )
    else:
        heading = "The ground was held, and it paid"
        opening = (
            f"{tilt_leader} took {tilt_value:.1f}% of the field tilt and turned it into "
            f"{leader_xg:.2f} expected goals against {trailer_xg:.2f}."
        )

    # A side with no final-third entries divides by zero here. It cannot happen
    # in a normal match and does happen in a partial parse, and the article
    # should still be written.
    survived = (
        f": {_ratio(100 * boxes, thirds):.0f}% of the entries survived the last "
        f"twenty metres" if thirds else ""
    )
    second = (
        f"The funnel is where it becomes concrete. {tilt_leader} reached the final third "
        f"{int(thirds)} times and the penalty area {int(boxes)}{survived}. "
        f"{tilt_trailer} arrived in the box {int(other_boxes)} times. Reaching the final "
        f"third is a function of having the ball; reaching the box is a function of "
        f"having somewhere to put it."
    )
    deep = m.of(tilt_leader, _num(m.hm, "deep_completions"),
                _num(m.am, "deep_completions"))
    prog = m.of(tilt_leader, _num(m.hm, "progressive_passes"),
                _num(m.am, "progressive_passes"))
    if leader_xg < trailer_xg:
        third = (
            f"{tilt_leader} played {int(prog)} progressive passes and completed "
            f"{int(deep)} into the deep attacking zone. The distinction those two "
            f"numbers draw is between moving the ball forward and moving it somewhere. "
            f"Progression that stops at the edge of the area pins an opponent back and "
            f"looks like control, and a settled block is perfectly content to concede "
            f"it — the defence is already where it wants to be. What actually breaks "
            f"one is a reception between the lines with a runner committed beyond the "
            f"last defender, and no amount of circulation in front of the block "
            f"substitutes for it."
        )
    else:
        third = (
            f"{int(prog)} progressive passes and {int(deep)} deep completions is "
            f"territory that arrived somewhere, which is the harder version to defend. "
            f"A side pinned into its own third can live with an opponent that "
            f"circulates in front of it; what it cannot live with is one that keeps "
            f"finding a body past the last line. The difference between the two rarely "
            f"shows in a possession figure and always shows in the box-entry count."
        )
    return Finding("territory", m.gap(home_tilt, away_tilt) + 0.15, Section(
        heading, [opening, second, third],
        m.visual("44_pitch_control.png", "24_dominating_zones.png"),
        pull_quote=f"{tilt_leader}: {int(thirds)} final-third entries → {int(boxes)} box entries",
    ))


def _finding_quality(m: _Match) -> Finding | None:
    """Shot volume against shot quality."""
    home_shots, away_shots = _num(m.hx, "shots"), _num(m.ax, "shots")
    home_per, away_per = _num(m.hx, "xG_per_shot"), _num(m.ax, "xG_per_shot")
    if not (home_shots or away_shots) or not (home_per or away_per):
        return None
    volume_leader, _v_trailer, volume_level = m.lead(home_shots, away_shots, tolerance=2)
    quality_leader, quality_trailer, quality_level = m.lead(
        home_per, away_per, tolerance=0.01)
    best, worst = (m.of(quality_leader, home_per, away_per),
                   m.of(quality_trailer, home_per, away_per))
    ratio = best / worst if worst else 0.0

    # A sentence that names a side must print that side's number first. Fixed
    # home-then-away order read as "Man City shot more often, 9 to 12" whenever
    # the away team led the count.
    volume = (
        f"the two sides took {int(home_shots)} and {int(away_shots)} shots"
        if volume_level else
        f"{volume_leader} shot more often, "
        f"{int(m.of(volume_leader, home_shots, away_shots))} to "
        f"{int(m.of(_v_trailer, home_shots, away_shots))}"
    )
    if quality_level:
        opening = (
            f"Neither the volume nor the value of the shooting separated them: "
            f"{volume}, at {best:.3f} and {worst:.3f} expected goals an attempt. Two "
            f"sides arriving at the same quality of chance by different routes is a "
            f"more interesting problem than one out-shooting the other, because it "
            f"means the difference in the match was made somewhere other than here."
        )
    else:
        # "On volume there was little in it" was printed even when one side had
        # out-shot the other by a third, contradicting the clause after the dash.
        lead_in = ("On volume there was little in it"
                   if volume_level else "One side shot more")
        opening = (
            f"{lead_in} — {volume}. On what those shots were "
            f"worth there was a great deal: {quality_leader} averaged {best:.3f} expected "
            f"goals a shot, {quality_trailer} {worst:.3f}."
            + (f" One side's average attempt was worth {ratio:.1f} times the other's."
               if ratio >= 1.5 else "")
        )
    home_ot, away_ot = _num(m.hx, "on_target"), _num(m.ax, "on_target")
    home_xgot, away_xgot = _num(m.hx, "xGoT"), _num(m.ax, "xGoT")
    second = (
        f"Shot count is the easiest attacking number to accumulate and the least "
        f"informative one to own. What survived to the goalkeeper is the better "
        f"question: {int(home_ot)} on target for {m.home} carrying {home_xgot:.2f} "
        f"post-shot xG, {int(away_ot)} for {m.away} carrying {away_xgot:.2f}. "
        f"Everything else was blocked, dragged wide, or taken from a position that "
        f"never justified the attempt."
    )
    # The heading counts the actual shots. Written as a fixed line it would have
    # said "nine good shots beat twelve bad ones" for every fixture.
    leader_shots = int(m.of(quality_leader, home_shots, away_shots))
    trailer_shots = int(m.of(quality_trailer, home_shots, away_shots))
    if quality_level:
        heading = "The shooting was the one thing they shared"
    elif leader_shots < trailer_shots:
        heading = (f"{_spell(leader_shots).capitalize()} good shots beat "
                   f"{_spell(trailer_shots)} lesser ones")
    else:
        heading = "Volume and quality pulled the same way"
    visuals = (m.team_visual("02_shot_map_{slug}.png", m.home)
               + m.team_visual("03_shot_map_{slug}.png", m.away)
               + m.visual("11_goalkeeper_saves.png"))
    home_big, away_big = _num(m.hx, "big_chances"), _num(m.ax, "big_chances")
    big_leader = m.of(quality_leader, home_big, away_big)
    big_trailer = m.of(quality_trailer, home_big, away_big)
    if big_leader > big_trailer:
        third = (
            f"Big chances split {int(big_leader)} to {int(big_trailer)} the same way, "
            f"which is the identical finding at a coarser resolution. A shot's value is "
            f"settled before it is struck: by the distance, the angle, the bodies "
            f"between ball and goal, and whether the goalkeeper had time to set. None "
            f"of that is the striker's doing. It is the pass before, and the movement "
            f"before that. A side taking low-value attempts is usually not shooting "
            f"badly — it is arriving badly, and the correction sits upstream of the shot."
        )
    else:
        third = (
            f"Big chances went the other way, {int(big_trailer)} to {int(big_leader)}, "
            f"which complicates the reading rather than settling it. Chance quality "
            f"averaged across every attempt and chance quality at the top end are "
            f"different questions: one side can hold the better mean while the other "
            f"gets the better single opportunities. The first describes how a team "
            f"arrives; the second describes what it did with the two or three moments "
            f"that were always going to decide the match."
        )
    weight = m.gap(home_per, away_per) if not quality_level else 0.12
    return Finding("quality", weight, Section(
        heading, [opening, second, third], visuals,
        pull_quote=f"{quality_leader} {best:.3f} xG per shot  ·  {quality_trailer} {worst:.3f}",
    ))


def _finding_game_state(m: _Match) -> Finding | None:
    """What the scoreline did to both sides."""
    states = ("leading", "drawing", "trailing")
    rows = {}
    for label, row in (("home", m.hm), ("away", m.am)):
        rows[label] = {s: (_num(row, f"game_state_{s}_xG"),
                           _num(row, f"game_state_{s}_shots"),
                           _num(row, f"game_state_{s}_box_entries")) for s in states}
    spread = max(
        abs(rows["home"][s][0] - rows["away"][s][0]) for s in states)
    if spread < 0.4:
        return None

    lines = []
    for team, key in ((m.home, "home"), (m.away, "away")):
        best_state = max(states, key=lambda s: rows[key][s][0])
        xg, shots, boxes = rows[key][best_state]
        if xg <= 0.05:
            continue
        lines.append(
            f"{team} produced {xg:.2f} of their expected goals while {best_state} "
            f"— {int(shots)} {_plural(shots, 'shot')} and {int(boxes)} box "
            f"{_plural(boxes, 'entry', 'entries')}"
        )
    if not lines:
        return None

    opening = (
        "Neither side played the same match for ninety minutes, and the scoreline is "
        "the reason. " + "; ".join(lines) + "."
    )
    second = (
        "This is the single most under-read board in a match report. A side that stops "
        "attacking once ahead and a side that only attacks once behind produce the same "
        "ninety-minute totals as two teams who played each other evenly, and the totals "
        "are what most reports quote. Split by state and the shape of the night appears: "
        "who chose the game, and who was made to play it."
    )
    if m.first_goal:
        when, scorer, _minute = m.first_goal
        moment = (
            f"The first goal arrived {when}"
            + (f", through {scorer}" if scorer else "")
            + ", and from that point the two teams were solving different problems. "
        )
    else:
        moment = "Until the first goal both sides were solving the same problem. "
    early = bool(m.first_goal and m.first_goal[2] <= 25)
    if early:
        third = (
            moment +
            "An opening goal that early does not decide a match, but it does decide "
            "what the next hour is going to look like. One side could drop its line and "
            "treat possession as time removed from the clock; the other had to push "
            "bodies past the ball and live with the space that leaves. Full-match "
            "averages then describe two different games added together, and any "
            "conclusion drawn from them without splitting by state is describing "
            "something that never happened."
        )
    else:
        third = (
            moment +
            "Because the score stayed level for so long, most of what both sides did "
            "was done under the same conditions — which makes the period after it "
            "changed the more interesting one to isolate. The totals are less "
            "misleading here than in a match settled early, but they still average "
            "across a break in the middle, and the split is where the reaction to it "
            "shows."
        )
    return Finding("game_state", 0.35 + spread / 4, Section(
        "The scoreline changed both teams", [opening, second, third],
        m.visual("33_game_state_splits.png", "35_match_momentum.png"),
    ))


def _finding_transition(m: _Match) -> Finding | None:
    """The match between the phases."""
    home_t, away_t = _num(m.hm, "transitions"), _num(m.am, "transitions")
    home_txg, away_txg = _num(m.hm, "transition_xG"), _num(m.am, "transition_xG")
    if not (home_t or away_t) or max(home_txg, away_txg) < 0.2:
        return None
    leader, trailer, level = m.lead(home_txg, away_txg, tolerance=0.15)
    leader_xg, trailer_xg = (m.of(leader, home_txg, away_txg),
                             m.of(trailer, home_txg, away_txg))
    leader_count = m.of(leader, home_t, away_t)
    trailer_count = m.of(trailer, home_t, away_t)
    leader_rate = m.of(leader, _num(m.hm, "regain_to_shot_rate"),
                       _num(m.am, "regain_to_shot_rate"))
    trailer_rate = m.of(trailer, _num(m.hm, "regain_to_shot_rate"),
                        _num(m.am, "regain_to_shot_rate"))

    if level:
        opening = (
            f"Broken play was shared and so was what came of it: {int(leader_count)} "
            f"transitions for {leader} and {int(trailer_count)} for {trailer}, worth "
            f"{leader_xg:.2f} and {trailer_xg:.2f} expected goals. When both sides "
            f"convert the open field at the same rate, the phase stops being the "
            f"explanation for anything and the difference has to be found elsewhere."
        )
    else:
        opening = (
            f"Broken play was close to shared — {int(leader_count)} transitions for "
            f"{leader}, {int(trailer_count)} for {trailer} — and the value taken from it "
            f"was not: {leader_xg:.2f} expected goals against {trailer_xg:.2f}."
        )
    # "The same asymmetry" was asserted whatever the two rates said. Arsenal
    # turned 4.4% of regains into a shot and Man City 4.5%, and the sentence
    # still claimed the transition gap ran through them.
    tail = (
        "The four seconds after winning the ball are when an opponent is at its most "
        "stretched, and a side that spends them securing possession has made a "
        "defensible choice and a costly one."
    )
    if abs(leader_rate - trailer_rate) <= 0.5:
        second = (
            f"It does not run through the regains, though: {leader} turned "
            f"{leader_rate:.1f}% of possession regains into a shot and {trailer} "
            f"{trailer_rate:.1f}%, which is the same rate. The advantage was in what the "
            f"transitions were worth, not in how often either side got a shot away from "
            f"one — so it was built by the positions the ball was won in and the runs "
            f"available from them, not by speed off the turnover."
        )
    elif leader_rate > trailer_rate:
        second = (
            f"The same asymmetry runs through the regains. {leader} turned "
            f"{leader_rate:.1f}% of possession regains into a shot; {trailer} managed "
            f"{trailer_rate:.1f}%. " + tail
        )
    else:
        second = (
            f"The regains cut the other way: {trailer} turned {trailer_rate:.1f}% of "
            f"possession regains into a shot against {leader}'s {leader_rate:.1f}%. "
            f"Getting a shot away more often from broken play while taking less value "
            f"out of it means the attempts themselves were worth less — the phase was "
            f"reached, the position at the end of it was not. " + tail
        )
    leader_vuln = m.of(leader, _num(m.hm, "rest_defence_vulnerability"),
                       _num(m.am, "rest_defence_vulnerability"))
    trailer_vuln = m.of(trailer, _num(m.hm, "rest_defence_vulnerability"),
                        _num(m.am, "rest_defence_vulnerability"))
    if leader_vuln > trailer_vuln:
        third = (
            f"The other half of transition is what happens when you lose the ball, and "
            f"here it cuts against the same side: {leader} were exposed on "
            f"{leader_vuln:.1f}% of their advanced losses against {trailer_vuln:.1f}%. "
            f"Taking more from broken play while giving more back is the profile of a "
            f"team that has chosen risk on both sides of the ball, which is a strategy "
            f"rather than an oversight — and one that needs the attacking half to keep "
            f"paying."
        )
    else:
        third = (
            f"{leader} also gave less back: exposed on {leader_vuln:.1f}% of their "
            f"advanced losses against {trailer}'s {trailer_vuln:.1f}%. Winning the "
            f"transition game at both ends usually says less about counter-attacking "
            f"and more about rest defence — where players stand relative to each other "
            f"at the moment possession turns, not how many of them are behind the ball. "
            f"A team can keep seven back and still concede the break if the distances "
            f"were wrong before it lost the ball."
        )
    return Finding("transition", m.gap(home_txg, away_txg) or 0.10, Section(
        "The open field" if level else "Who used the broken play",
        [opening, second, third],
        m.visual("32_transition_outcomes.png", "49_press_triggers.png"),
        pull_quote=f"Transition xG: {leader} {leader_xg:.2f}  ·  {trailer} {trailer_xg:.2f}",
    ))


def _finding_press(m: _Match) -> Finding | None:
    """Pressing, and whether it was worth anything."""
    # PPDA is not in every export. Defaulting it to zero printed "PPDA read
    # 0.00 for Arsenal and 0.00 for Man City" — a figure that cannot occur,
    # since a side allowing no passes per defensive action has not played.
    home_ppda, away_ppda = _num(m.hm, "ppda", 0.0), _num(m.am, "ppda", 0.0)
    ppda = (f" PPDA read {home_ppda:.2f} for {m.home} and {away_ppda:.2f} for "
            f"{m.away}, but that number only describes how hard a side pressed, "
            f"never what it collected." if home_ppda > 0 and away_ppda > 0 else "")
    home_hr, away_hr = _num(m.hm, "high_regains"), _num(m.am, "high_regains")
    if not (home_hr or away_hr):
        return None
    presser, other, level = m.lead(home_hr, away_hr, tolerance=3)
    presser_hr, other_hr = m.of(presser, home_hr, away_hr), m.of(other, home_hr, away_hr)
    presser_cp = m.of(presser, _num(m.hm, "counterpress_success_rate"),
                      _num(m.am, "counterpress_success_rate"))
    other_cp = m.of(other, _num(m.hm, "counterpress_success_rate"),
                    _num(m.am, "counterpress_success_rate"))
    if level:
        opening = (
            f"Both sides pressed, and both got roughly the same for it: "
            f"{int(presser_hr)} high regains against {int(other_hr)}, with "
            f"counterpressing succeeding {presser_cp:.1f}% and {other_cp:.1f}% of the "
            f"time. Two teams squeezing each other equally hard tends to produce a "
            f"match decided in the moments the press does not reach."
        )
    else:
        opening = (
            f"{presser} won the ball back in the opponent's territory {int(presser_hr)} "
            f"times to {int(other_hr)}, and recovered it within seconds of losing it "
            f"{presser_cp:.1f}% of the time against {other_cp:.1f}%."
        )
    second = (
        "Pressing metrics describe the press, not the point of it. The question that "
        "matters is what the regains became, and a side can lead every pressing board "
        "on the page while converting almost none of the disorder it created."
    )
    presser_rate = m.of(presser, _num(m.hm, "regain_to_shot_rate"),
                        _num(m.am, "regain_to_shot_rate"))
    other_rate = m.of(other, _num(m.hm, "regain_to_shot_rate"),
                      _num(m.am, "regain_to_shot_rate"))
    # A tenth of a percentage point is a tie, and calling it "paid for" made
    # the paragraph contradict the one above it, which had just said a side can
    # lead every pressing board while converting none of the disorder.
    if abs(presser_rate - other_rate) <= 0.5:
        third = (
            f"It bought nothing either way: {presser_rate:.1f}% of {presser}'s regains "
            f"became a shot and {other_rate:.1f}% of {other}'s, which is the same rate "
            f"from a different number of regains." + ppda + " Pressing harder and "
            f"converting at the opponent's rate means the extra regains arrived in "
            f"positions no more dangerous than the ordinary ones — the press was won "
            f"where it did not matter."
        )
    elif presser_rate > other_rate:
        third = (
            f"And it was paid for: {presser_rate:.1f}% of regains became a shot against "
            f"{other_rate:.1f}%." + ppda + " Winning the ball high and turning "
            f"it into an attempt in the seconds that follow is the whole return on the "
            f"cost, and it is the half of the press most often left unmeasured."
        )
    else:
        third = (
            f"It was not paid for. {presser} turned {presser_rate:.1f}% of regains into "
            f"a shot; {other} managed {other_rate:.1f}% from fewer of them." + ppda +
            f" The four seconds after a high regain are when an opponent "
            f"is at its most stretched and has nobody arranged between the ball and the "
            f"goal; a side that wins it there and takes a touch to settle has spent the "
            f"advantage it just spent ninety minutes manufacturing."
        )
    return Finding("press", (m.gap(home_hr, away_hr) * 0.8) or 0.10, Section(
        "Two presses, cancelling out" if level else "The press, and what it was worth",
        [opening, second, third],
        m.visual("31_ppda_pressing.png")
        + m.team_visual("27_high_regains_{slug}.png", presser),
    ))


def _finding_player(m: _Match) -> Finding | None:
    """The individual the match ran through."""
    best = m.top_player()
    if best is None:
        return None
    name = _text(best.get("player"))
    if not name:
        return None
    team = _text(best.get("team"))
    xt = _number(best.get("sequence_xT"))
    chain = _number(best.get("xGChain"))
    sequences = _whole(best.get("sequences"))
    if xt <= 0:
        return None

    rival_team = m.away if team == m.home else m.home
    rival = m.top_player(rival_team)
    comparison = ""
    if rival is not None:
        rival_name = _text(rival.get("player"))
        rival_xt = _number(rival.get("sequence_xT"))
        if rival_name and rival_xt > 0:
            comparison = (
                f" {rival_name} led {rival_team} on {rival_xt:.2f}, which is the "
                f"comparison worth making: not who touched the ball most, but whose "
                f"involvement kept arriving in dangerous possessions."
            )

    opening = (
        f"{name} was involved in more of the valuable attacking play than anyone on the "
        f"pitch: {xt:.2f} sequence expected threat from {sequences} possessions, and "
        f"{chain:.2f} xGChain.{comparison}"
    )
    second = (
        "Sequence threat credits every player in a possession that added danger, not "
        "only the one who finished it. It is the closest single number to the question "
        "a coach actually asks after a match — who moved this team forward — "
        "and it is deliberately blind to whether the last touch went in."
    )
    buildup = _number(best.get("xGBuildup"))
    buildup_share = buildup / chain if chain else 0.0
    if buildup_share >= 0.5:
        third = (
            f"Most of it — {buildup:.2f} xGBuildup of {chain:.2f} xGChain — survives "
            f"even after the shot and the pass before it are stripped out. That is the "
            f"profile of a player who builds the danger rather than finishes it: the "
            f"reception that turned a phase forward, the pass that moved a defender, "
            f"the run that dragged a marker out of the space somebody else used. Goals "
            f"and assists measure the two touches nearest the finish and miss all of it."
        )
    else:
        third = (
            f"Only {buildup:.2f} of the {chain:.2f} xGChain survives once the shot and "
            f"the pass before it are removed, which places the contribution close to "
            f"the finish rather than in the construction. That is not a lesser role, "
            f"but it is a different one, and it makes the player more dependent on the "
            f"phase behind him: take away the supply and this profile has less to do."
        )
    return Finding("player", 0.30, Section(
        f"{name} carried it", [opening, second, third],
        m.visual("34_player_sequence_leaders.png", "43_action_value.png"),
    ))


def _finding_profiles(m: _Match, events) -> Finding | None:
    """The three players each side leaned on, as radars.

    The report's appendix carries one for every participant. An article wants
    the few the match actually turned on, which is a different selection and
    the reason this asks the ranking rather than listing the folder.
    """
    try:
        from player_radar import top_players_per_team

        ranking = top_players_per_team(events, {
            "home_id": m.hm.get("team_id"), "away_id": m.am.get("team_id"),
            "home_name": m.home, "away_name": m.away,
        }, n=RADARS_PER_TEAM)
    except Exception:
        return None

    visuals, named = [], {}
    for side, team in (("home", m.home), ("away", m.away)):
        picked = []
        for player in ranking.get(side, [])[:RADARS_PER_TEAM]:
            radar = m.radar(team, player)
            if radar is not None:
                visuals.append(radar)
                picked.append(str(player))
        if picked:
            named[team] = picked
    if not visuals:
        return None

    lines = [f"{team}: {', '.join(players)}" for team, players in named.items()]
    opening = (
        f"{_spell(len(visuals)).capitalize()} profiles, the "
        f"{_spell(RADARS_PER_TEAM)} each side leaned on most. "
        + ". ".join(lines) + "."
    )
    second = (
        "Each radar is one match, not a rating. The bars are percentiles against every "
        "player on the pitch that afternoon, so a full wedge means the player led this "
        "fixture on that action and nothing more; the chip beside it carries the raw "
        "number, because a percentile with no count behind it can dress two touches up "
        "as dominance."
    )
    third = (
        "Read them for shape rather than area. A defender with an empty attacking half "
        "and a full defensive one has done the job asked of him, and the interesting "
        "players are the ones whose profile does not match the position they were "
        "listed in — the full-back with a creator's passing segments, the forward whose "
        "value sits in build-up rather than finishing."
    )
    return Finding("profiles", 0.22,
                   Section("The players it turned on", [opening, second, third],
                           visuals, gallery=True))


BUILDERS = (
    _finding_territory,
    _finding_quality,
    _finding_game_state,
    _finding_transition,
    _finding_press,
    _finding_player,
)


# --------------------------------------------------------------------------
# assembly
# --------------------------------------------------------------------------

def _title(m: _Match) -> tuple[str, str]:
    """The headline, taken from whatever actually separated the two sides.

    There used to be three titles in the whole file — one per result shape — so
    every ordinary win was published as "How X Took Y Apart" and the headline
    said nothing a reader could not see in the scoreline.

    Each candidate below owns a condition and a strength. The strongest
    condition writes the headline, which means two matches share a title only
    when the same thing decided them by the same margin. It is varied because
    matches differ, not because anything is random: the same fixture always
    produces the same headline.
    """
    candidates: list[tuple[float, str, str]] = []

    def offer(strength: float, headline: str, standfirst: str) -> None:
        if strength > 0:
            candidates.append((strength, headline, standfirst))

    result = f"{m.winner} beat {m.loser} {m.score}" if m.winner else \
             f"{m.home} and {m.away} finished {m.score}"

    # -- the result against the process ----------------------------------
    xg_leader, _xg_trailer, xg_level = m.lead(m.home_xg, m.away_xg, tolerance=0.15)
    if m.winner and not xg_level and xg_leader != m.winner:
        offer(3.0,
              f"{m.loser} Won Everything But The Match",
              f"{result}. The expected goals finished "
              f"{m.of(m.loser, m.home_xg, m.away_xg):.2f} to "
              f"{m.of(m.winner, m.home_xg, m.away_xg):.2f} the other way.")

    # -- territory that never became threat -------------------------------
    home_tilt, away_tilt = _num(m.hm, "field_tilt"), _num(m.am, "field_tilt")
    tilt_leader, tilt_trailer, tilt_level = m.lead(home_tilt, away_tilt, tolerance=6.0)
    tilt_xg = m.of(tilt_leader, m.home_xg, m.away_xg)
    other_xg = m.of(tilt_trailer, m.home_xg, m.away_xg)
    if not tilt_level and other_xg > tilt_xg:
        offer(1.2 + m.gap(home_tilt, away_tilt),
              f"{tilt_leader} Had The Ball. {tilt_trailer} Had The Chances.",
              f"{result}. {tilt_leader} held "
              f"{m.of(tilt_leader, home_tilt, away_tilt):.1f}% of the final-third "
              f"passing and {tilt_xg:.2f} expected goals for it.")

    # -- the last twenty metres ------------------------------------------
    def survival(side_row):
        thirds = _num(side_row, "final_third_entries")
        return _ratio(100 * _num(side_row, "box_entries"), thirds) if thirds else 0.0

    home_survival, away_survival = survival(m.hm), survival(m.am)
    if min(home_survival, away_survival) > 0:
        better, worse, survival_level = m.lead(home_survival, away_survival, tolerance=8.0)
        if not survival_level:
            offer(1.0 + m.gap(home_survival, away_survival),
                  "The Last Twenty Metres Decided It",
                  f"{result}. {better} turned "
                  f"{m.of(better, home_survival, away_survival):.0f}% of its "
                  f"final-third entries into the box; {worse} "
                  f"{m.of(worse, home_survival, away_survival):.0f}%.")

    # -- what a shot was worth -------------------------------------------
    home_per, away_per = _num(m.hx, "xG_per_shot"), _num(m.ax, "xG_per_shot")
    home_shots, away_shots = _num(m.hx, "shots"), _num(m.ax, "shots")
    quality, blunt, quality_level = m.lead(home_per, away_per, tolerance=0.02)
    if not quality_level and m.of(quality, home_shots, away_shots) < m.of(blunt, home_shots, away_shots):
        offer(1.1 + m.gap(home_per, away_per),
              f"{quality} Shot Less And Meant It More",
              f"{result}. {quality} averaged "
              f"{m.of(quality, home_per, away_per):.3f} expected goals an attempt "
              f"against {m.of(blunt, home_per, away_per):.3f}, from fewer shots.")

    # -- broken play ------------------------------------------------------
    home_txg, away_txg = _num(m.hm, "transition_xG"), _num(m.am, "transition_xG")
    if max(home_txg, away_txg) >= 0.5:
        breaker, held, transition_level = m.lead(home_txg, away_txg, tolerance=0.25)
        if not transition_level:
            offer(0.9 + m.gap(home_txg, away_txg),
                  f"{breaker} Won The Match In The Broken Moments",
                  f"{result}. Transitions were worth "
                  f"{m.of(breaker, home_txg, away_txg):.2f} expected goals to "
                  f"{breaker} and {m.of(held, home_txg, away_txg):.2f} to {held}.")

    # -- an early goal that removed the level phase ------------------------
    if m.first_goal and m.first_goal[2] <= 2 and m.winner:
        when, scorer, _minute = m.first_goal
        offer(1.3,
              "The Match Was Framed Before It Started",
              f"{result}. The opening goal arrived {when}"
              + (f", through {scorer}" if scorer else "")
              + ", so neither side ever played a level minute.")

    # -- finishing well away from the chances ------------------------------
    scored, expected = m.home_goals + m.away_goals, m.home_xg + m.away_xg
    if scored >= expected + 1.5:
        offer(0.8 + (scored - expected) / 6,
              "The Finishing Outran The Football",
              f"{result}. {scored} goals came from {expected:.2f} combined "
              f"expected goals, so conversion, not creation, set the margin.")
    elif expected >= scored + 1.5:
        offer(0.8 + (expected - scored) / 6,
              "The Chances Were There. The Finishing Was Not.",
              f"{result}. {expected:.2f} combined expected goals produced "
              f"{scored}, so the scoreline understates what was built.")

    # -- a one-sided scoreline --------------------------------------------
    if m.winner and abs(m.home_goals - m.away_goals) >= 3:
        offer(0.85,
              f"{m.winner} Made It Look Simple",
              f"{result}, and the margin was not an accident. This is where it "
              f"came from.")

    # -- the fallbacks, which only win when nothing above applies ----------
    if m.winner:
        offer(0.5,
              f"How {m.winner} Took {m.loser} Apart",
              f"{result}, and the underlying numbers agree with the scoreline. "
              f"This is how the margin was built.")
    else:
        offer(0.6,
              "The Draw That Was Not Even",
              f"{result}. Almost nothing else about the match was level.")
    offer(0.1, f"{m.home} vs {m.away}, Read From The Data",
          f"{result}. What the match record says about how it happened.")

    strength, headline, standfirst = max(candidates, key=lambda row: row[0])
    return headline, standfirst


def _cover_image(out_dir: Path) -> Path | None:
    """The report's own cover page, rendered for the article to open on.

    Not the bare artwork: the cover is the artwork plus the badge, both crests,
    the score, the two-colour rule and the finding underneath, and the article
    should open on the same page the report does. Rendering page one of the
    finished PDF makes them identical by construction rather than by two
    layouts agreeing.
    """
    pdf = out_dir / "full_visual_redesign_real_data.pdf"
    target = out_dir / "article_cover.png"
    if pdf.exists():
        try:
            import fitz

            document = fitz.open(pdf)
            try:
                page = document[0]
                # 200dpi: sharp in Word and on a phone, without a 20MB page.
                page.get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72)).save(target)
            finally:
                document.close()
            if target.exists():
                return target
        except Exception:
            pass
    # No report yet, or no renderer: the artwork alone is better than nothing.
    artwork = out_dir / "cover_art.png"
    return artwork if artwork.exists() else None


def _closing(m: _Match, used: list[str]) -> Section:
    leader, _trailer, level = m.lead(m.home_xg, m.away_xg, tolerance=0.15)
    if m.winner and not level and leader != m.winner:
        first = (
            f"None of this makes {m.winner}'s win undeserved in any sense that matters "
            f"to a league table. Results are the currency and they took the points. But "
            f"the performance underneath a result is the part that carries into the next "
            f"match, and on the evidence here it belongs to {m.loser}."
        )
    elif m.winner:
        first = (
            f"{m.winner} won this the way the numbers say they should have. The margin "
            f"was built in the phases above rather than at the end of them, which is the "
            f"version of a win most likely to repeat."
        )
    else:
        first = (
            f"A draw hides more than any other result. On this evidence the two sides "
            f"were not the same team for ninety minutes, whatever the scoreline says."
        )
    gap = m.gap(m.home_xg, m.away_xg)
    if gap >= 0.35:
        second = (
            "One match is never a verdict, but a gap this wide is harder to explain away "
            "than a narrow one. Conversion swings hardest over a handful of shots; the "
            "distance between what the two sides created does not swing nearly as far. "
            "What is worth carrying forward is the mechanism — where the chances came "
            "from, which phase produced them, and whether it survives a different "
            "opponent."
        )
    else:
        second = (
            "One match is never a verdict, and this one was closer underneath than the "
            "scoreline suggests. Over ninety minutes conversion moves further than "
            "anything else on the page, so a margin built this narrowly should be held "
            "loosely. The mechanism is the part worth watching again: which phase "
            "produced the chances, and whether either side can reproduce it."
        )
    return Section("What to take from it", [first, second],
                   m.visual("14_post_match_advanced_dashboard.png"))


_GALLERY_ORDER = (
    "04_goals_breakdown", "40_win_probability", "45_sequence_types",
    "05a_pass_network", "05b_pass_network", "06a_pass_network", "06b_pass_network",
    "22a_average_positions", "22b_average_positions",
    "23a_average_positions", "23b_average_positions",
    "07_xt_map", "08_xt_map", "09_pass_map", "10_pass_map",
    "29_pass_targets", "30_pass_targets", "12_zone14", "13_zone14",
    "47_unlocking", "48_unlocking", "16_progressive", "17_progressive",
    "41_playing_through", "42_playing_through", "18_crosses", "19_crosses",
    "25_box_entries", "26_box_entries", "36_set_pieces", "15_xt_per_minute",
    "20_defensive_activity", "21_defensive_activity", "39_defensive_shape",
    "27_high_regains", "28_high_regains", "37_ball_losses", "38_ball_losses",
)


def _gallery(m: _Match, used: set[str]) -> Section | None:
    """Everything the argument did not need, in a reading order.

    The article picks its evidence, but the package produced fifty-three
    visuals and a reader who wants the rest should not have to open the PDF to
    find them. They go at the end, after the argument has been made.
    """
    remaining = []
    for prefix in _GALLERY_ORDER:
        for path in sorted(m.out.glob(f"{prefix}*.png")):
            if path.name not in used:
                remaining.append(path)
                used.add(path.name)
    # Then everything the order does not name. Without this a visual reached the
    # article only if it was either chosen as evidence or listed above, and
    # 31_ppda_pressing.png was neither whenever the press finding was trimmed.
    for path in sorted(m.out.glob("[0-9]*.png")):
        if path.name not in used:
            remaining.append(path)
            used.add(path.name)
    if not remaining:
        return None
    opening = (
        f"The argument above used the boards that carried it. These are the rest, in "
        f"the order a full read would take them — how each side built possession, "
        f"where it tried to progress, what it did in the final third and what it did "
        f"without the ball — each with the reading that belongs to it."
    )
    return Section("The rest of the evidence", [opening], remaining, gallery=True)


def build_article(
    events: pd.DataFrame,
    xg: pd.DataFrame,
    team_metrics: pd.DataFrame,
    player_metrics: pd.DataFrame,
    match_info: dict,
    out_dir: Path | str,
    *,
    max_sections: int = 5,
) -> Article:
    """Derive one article from a fixture's frames."""
    m = _Match(events, xg, team_metrics, player_metrics, match_info, out_dir)

    opener = _finding_result(m)
    rest = [f for f in (builder(m) for builder in BUILDERS) if f is not None]
    profiles = _finding_profiles(m, events)
    if profiles is not None:
        rest.append(profiles)
    rest.sort(key=lambda f: f.weight, reverse=True)

    # Take the strongest findings, then keep taking while the piece is short of
    # the length it was commissioned at. An even match has weaker findings, not
    # fewer things worth saying, and stopping at a fixed count left one at 773
    # words against a 1200 floor.
    # Every finding, strongest first. There is no ceiling to trim against, and
    # a finding the match actually supports is not worth dropping to save
    # words the reader was never promised.
    del max_sections
    chosen = ([opener] if opener else []) + rest

    title, standfirst = _title(m)
    strap = " · ".join(
        part for part in (m.competition.upper(),
                          f"{m.home.upper()} {m.score} {m.away.upper()}") if part)
    # The profiles section is the one finding that always earns its place: an
    # article about a match that never shows a player is missing the people.
    if profiles is not None and profiles not in chosen:
        chosen.append(profiles)

    closing = _closing(m, [f.key for f in chosen])

    sections = [f.section for f in chosen] + [closing]
    used = {Path(v).name for s in sections for v in s.visuals}
    gallery = _gallery(m, used)
    if gallery is not None:
        sections.append(gallery)
    cover = _cover_image(m.out)
    context = None
    try:
        from tactical_pdf_report import build_context

        context = build_context(events, xg, team_metrics, player_metrics, match_info)
    except Exception:
        context = None
    return Article(title, standfirst, strap, sections,
                   cover, m.home, m.away, context)


# --------------------------------------------------------------------------
# Word output
# --------------------------------------------------------------------------

def render_docx(article: Article, path: Path | str,
                home: str = "", away: str = "") -> Path:
    """Write the article as a .docx built for pasting straight into Substack.

    Real Heading 1/2 styles, because that is what an editor's paste reads;
    images at a fixed measure with a caption under each; no tables, which
    Substack breaks.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    path = Path(path)
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.45

    if article.cover is not None and Path(article.cover).exists():
        cover = document.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover.add_run().add_picture(str(article.cover), width=Inches(6.4))

    strap = document.add_paragraph()
    run = strap.add_run(article.strap)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    document.add_heading(article.title, level=1)

    stand = document.add_paragraph()
    stand_run = stand.add_run(article.standfirst)
    stand_run.italic = True
    stand_run.font.size = Pt(13.5)
    stand_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for section in article.sections:
        document.add_heading(section.heading, level=2)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
        if section.pull_quote:
            quote = document.add_paragraph()
            quote.paragraph_format.left_indent = Inches(0.35)
            quote_run = quote.add_run(section.pull_quote)
            quote_run.bold = True
            quote_run.font.size = Pt(13)
        for visual in section.visuals:
            holder = document.add_paragraph()
            holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            holder.add_run().add_picture(str(visual), width=Inches(6.2))
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(_caption(visual, home, away))
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            caption_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            # In the appendix a caption alone is a label. Each board gets the
            # reading the report writes under it, from the same source, so the
            # two documents cannot say different things about the same picture.
            note = _analysis(visual, article.context)
            if note:
                body = document.add_paragraph()
                body_run = body.add_run(note)
                body_run.font.size = Pt(10.5)
                body_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "All figures derived from WhoScored/Opta event data. Visuals generated with "
        "an open-source pipeline.")
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


# What each board actually shows. Every visual the package produces has an
# entry: the fallback turned a filename into title case, so thirty-six of the
# fifty-three carried "Xt Map Arsenal" where a caption belonged.
#
# Keys are matched against the filename stem, longest first, and {team} is
# filled from the side named in it.
_CAPTIONS = {
    "xg_flow": "Cumulative expected goals. Each step is a chance; the flat stretches are possession that never became one.",
    "goals_breakdown": "Every goal with the pass that made it and the minute it arrived.",
    "goal_origins": "Each goal traced back to the moment its possession began.",
    "win_probability": "How the likely result moved as chances arrived.",
    "xt_per_minute": "Threat added minute by minute, both sides on one axis.",
    "match_momentum": "Expected-goal difference in five-minute windows.",
    "sequence_types": "How each side's danger was built: sustained possession, transition, or a restart.",
    "game_state_splits": "Output by scoreline state — leading, level, and chasing.",

    "pitch_control": "Distance-decayed influence: who held which space.",
    "dominating_zones": "Touch difference by zone. Positive is the home side.",
    "post_match_advanced_dashboard": "The match in thirty-two indicators.",

    "shot_map": "{team}'s shots, sized by chance quality.",
    "goalkeeper_saves": "What each goalkeeper had to deal with, on the goal frame.",
    "set_pieces": "What the restarts produced for each side.",

    "pass_network_{half}": "{team}'s passing network, {half_word}. Node size is touches; line weight is passes between the pair.",
    "average_positions_{half}": "{team}'s average positions, {half_word}.",
    "pass_network": "{team}'s passing network. Node size is touches; line weight is passes between the pair.",
    "average_positions": "{team}'s average positions.",
    "pass_map": "Every pass {team} attempted, completed and not.",
    "pass_targets": "Where {team} aimed its passes — the density of intended receivers.",
    "xt_map": "Where {team}'s possession added threat.",

    "progressive": "{team}'s passes that moved play meaningfully forward.",
    "playing_through": "{team}'s passes that broke a defensive line.",
    "zone14": "{team} in zone 14 and the five vertical lanes.",
    "unlocking": "{team}'s receptions in the pocket ahead of the defensive line.",
    "box_entries": "How {team} entered the penalty area.",
    "crosses": "{team}'s delivery from wide, completed and not.",

    "defensive_activity": "Where {team} did its defensive work.",
    "defensive_shape": "The shape each side held out of possession.",
    "high_regains": "{team}'s regains in the opponent's territory.",
    "ppda_pressing": "Opponent passes allowed per defensive action. Lower presses harder.",
    "press_triggers": "What the opponent was doing when the ball was won high.",
    "ball_losses": "Where {team} gave the ball away, and which losses were punished.",
    "transition_outcomes": "What broken play produced for each side.",

    "player_sequence_leaders": "Involvement in valuable attacking sequences.",
    "action_value": "Every action priced in goals.",
}

_HALF_WORDS = {"1h": "first half", "2h": "second half"}


def _analysis(path, context) -> str:
    """The report's own reading of one visual, or "" if it has none."""
    if not context:
        return ""
    try:
        from tactical_pdf_report import visual_explanation

        return (visual_explanation(Path(path), context) or "").strip()
    except Exception:
        return ""


def _caption(path, home: str = "", away: str = "") -> str:
    """The caption for one visual, with its team and half filled in.

    ``home`` and ``away`` are needed because the filename carries a slug —
    "man_city" — and a caption should say "Man City".
    """
    stem = Path(path).stem
    lowered = stem.lower()

    if "player_radars" in str(path):
        return f"{Path(path).stem.replace('_', ' ')} — one match, percentiles against every player on the pitch."

    half = next((h for h in ("1h", "2h") if lowered.endswith(h)), "")

    # Match the slug as the trailing token, longest first. A substring search
    # gave "milan" the boards belonging to "inter_milan", and a side called
    # Cross every board whose name contains "crosses".
    body = lowered[: -(len(half) + 1)] if half else lowered
    team = ""
    for name in sorted((home, away), key=lambda n: len(_slug(n or "")), reverse=True):
        slug = _slug(name or "")
        if name and slug and (body == slug or body.endswith("_" + slug)):
            team = name
            break

    # Longest key first, so "pass_network_{half}" wins over "pass_network".
    # The half is tested separately from the prefix: the filename puts the team
    # slug between them ("05b_pass_network_arsenal_2h"), so a single contiguous
    # probe never matched and every half-specific caption fell through.
    for key in sorted(_CAPTIONS, key=len, reverse=True):
        if "{half}" in key:
            if not half:
                continue
            probe = key.replace("_{half}", "")
        else:
            probe = key
        if probe in lowered:
            text = _CAPTIONS[key]
            return (text.replace("{team}", team or "The side")
                        .replace("{half_word}", _HALF_WORDS.get(half, ""))
                        .replace("{half}", half))
    return stem.split("_", 1)[-1].replace("_", " ").title()


def build_match_article(
    events, xg, team_metrics, player_metrics, match_info, out_dir,
    players=None, *, strict: bool = True,
) -> Path | None:
    """Build the article and write it beside the package. None on failure.

    Refuses to write when the fixture does not hold together. Everything else
    in the pipeline reads the frames and writes about them with total
    confidence; an article is the artefact most likely to be published without
    a second look, so it is the one that should decline rather than describe a
    match that never happened.
    """
    try:
        if strict and players is not None:
            from match_sanity import describe, inspect

            problems = inspect(events, players, xg, match_info)
            if problems:
                print("  ! article not written — the fixture does not hold together:")
                print(describe(problems))
                return None

        article = build_article(events, xg, team_metrics, player_metrics,
                                match_info, out_dir)
        target = Path(out_dir) / "match_article.docx"
        return render_docx(article, target, article.home, article.away)
    except Exception as error:
        # The article is optional — a failure here must not take the package
        # down with it — but a silent None is indistinguishable from a refusal,
        # and debugging one costs an hour. Say what broke.
        print(f"  ! article not written — {type(error).__name__}: {error}")
        return None
